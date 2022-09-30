import docx
from docx.shared import Cm
from docx.oxml.ns import qn

from win32com.client import gencache
from win32com.client import constants, gencache

doc = docx.Document()


doc.add_heading('标题', 0)
doc.add_heading('2022-09-30', 1)

doc.styles['Normal'].font.name = u'华文行楷'
doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


doc.add_paragraph('尊敬的 ** 先生/女士：')

# 设置缩进
# paragraph = doc.add_paragraph("42090risafjd;a fafsafsa来dddddddddddddddddddddddddddddddddddddddddddd")

paragraph = doc.add_paragraph()

paragraph_format = paragraph.paragraph_format
# 设置首行缩进
paragraph_format.first_line_indent = Cm(1)

run = paragraph.add_run("dddddddddddddddddddd0324您好，明天就要放假了")

run.bold = True


doc.save('D://invite6.docx')


# word转PDF
word = gencache.EnsureDispatch('Word.Application')
d = word.Documents.Open('D://invite6.docx', ReadOnly=1)
d.ExportAsFixedFormat('D://invite6.pdf', constants.wdExportFormatPDF, Item=constants.wdExportDocumentWithMarkup,
                      CreateBookmarks=constants.wdExportCreateHeadingBookmarks)

word.Quit(constants.wdDoNotSaveChanges)
